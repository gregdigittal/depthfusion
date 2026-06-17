"""Document parser for the ingestion framework (E-53).

Supports:
  - ``.docx`` — via python-docx
  - ``.pdf``  — via pypdf
  - ``.txt``  — plain text UTF-8
  - ``.md``   — Markdown (treated as plain text)

Parse budget (T-599):
  Set ``DEPTHFUSION_PARSE_MAX_BYTES`` to a byte-count limit.  Documents
  exceeding the limit are quarantined instead of parsed.  The default is
  ``50 * 1024 * 1024`` (50 MiB).  Set to ``0`` to disable (no limit).

  Quarantine destination: by default, quarantine metadata is written as
  JSON files to the ``data/quarantine/`` directory (relative to the
  current working directory).  Pass a custom
  :class:`~depthfusion.parsers.documents.base.QuarantineStore` instance
  to override the destination.

Usage::

    from depthfusion.ingest.parser import DocumentParser

    parser = DocumentParser()
    doc = parser.parse("/path/to/report.pdf", "application/pdf")
    print(doc.text[:200])
"""
from __future__ import annotations

import io
import json
import logging
import os
import pathlib
import re
import threading
from typing import TYPE_CHECKING

from depthfusion.ingest.models import ParsedDocument

if TYPE_CHECKING:
    from depthfusion.parsers.documents.base import QuarantineEntry

_logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Parse-budget helpers (T-599)
# ---------------------------------------------------------------------------

#: Default parse-budget: 50 MiB (goal spec).
_DEFAULT_PARSE_MAX_BYTES: int = 50 * 1024 * 1024

#: Default quarantine directory (relative to CWD).
_DEFAULT_QUARANTINE_DIR: str = "data/quarantine"


def _get_parse_max_bytes() -> int:
    """Return the active parse budget from env or the built-in default.

    ``DEPTHFUSION_PARSE_MAX_BYTES=0`` disables the budget (no limit).
    """
    raw = os.environ.get("DEPTHFUSION_PARSE_MAX_BYTES", "")
    if raw.strip():
        try:
            return int(raw.strip())
        except ValueError:
            pass
    return _DEFAULT_PARSE_MAX_BYTES


# ---------------------------------------------------------------------------
# File-based quarantine store (T-599)
# ---------------------------------------------------------------------------

class FileQuarantineStore:
    """Persist quarantine entries as JSON files under a directory.

    Each quarantined document produces one ``<safe-name>.json`` file in
    *quarantine_dir*.  The file records the source_id, error reason,
    timestamp, and raw size.

    Thread-safe via an internal lock.

    Args:
        quarantine_dir: Directory path for quarantine files.  Created on
                        demand if it does not exist.  Defaults to
                        ``data/quarantine`` relative to the current
                        working directory.
    """

    def __init__(self, quarantine_dir: str | pathlib.Path | None = None) -> None:
        self._dir = pathlib.Path(quarantine_dir or _DEFAULT_QUARANTINE_DIR)
        self._lock = threading.Lock()

    def add(self, entry: object) -> None:  # QuarantineEntry from base module
        """Write *entry* as a JSON file in the quarantine directory."""
        from depthfusion.parsers.documents.base import QuarantineEntry

        if not isinstance(entry, QuarantineEntry):
            raise TypeError(f"Expected QuarantineEntry, got {type(entry)}")

        with self._lock:
            self._dir.mkdir(parents=True, exist_ok=True)
            # Build a filesystem-safe filename from source_id
            safe = re.sub(r"[^A-Za-z0-9._-]", "_", entry.source_id)[:200]
            dest = self._dir / f"{safe}.json"
            payload = {
                "source_id": entry.source_id,
                "error_message": entry.error_message,
                "timestamp": entry.timestamp,
                "raw_size_bytes": entry.raw_size_bytes,
                "retry_count": entry.retry_count,
                "max_retries": entry.max_retries,
                "next_retry_at": entry.next_retry_at,
                "last_error": entry.last_error,
            }
            _logger.warning(
                "Quarantined %s (%d bytes): %s",
                entry.source_id,
                entry.raw_size_bytes,
                entry.error_message,
            )
            dest.write_text(json.dumps(payload, indent=2), encoding="utf-8")

    def list_all(self) -> list[QuarantineEntry]:
        """Return all quarantine entries loaded from disk."""
        from depthfusion.parsers.documents.base import QuarantineEntry

        with self._lock:
            if not self._dir.exists():
                return []
            entries = []
            for f in self._dir.glob("*.json"):
                try:
                    data = json.loads(f.read_text(encoding="utf-8"))
                    entries.append(
                        QuarantineEntry(
                            source_id=data["source_id"],
                            error_message=data.get("error_message", ""),
                            timestamp=data.get("timestamp", ""),
                            raw_size_bytes=int(data.get("raw_size_bytes", 0)),
                            retry_count=int(data.get("retry_count", 0)),
                            max_retries=int(data.get("max_retries", 3)),
                            next_retry_at=data.get("next_retry_at", ""),
                            last_error=data.get("last_error", ""),
                        )
                    )
                except Exception:  # noqa: BLE001
                    pass
            return entries

    def get_json_path(self, source_id: str) -> pathlib.Path:
        """Return the path where *source_id* would be stored."""
        safe = re.sub(r"[^A-Za-z0-9._-]", "_", source_id)[:200]
        return self._dir / f"{safe}.json"


# Sentinel for optional imports — only loaded if the caller actually uses
# that format so the package stays importable even without python-docx /
# pypdf installed.
_docx_available: bool = False
_pypdf_available: bool = False

try:
    import docx as _docx  # python-docx
    _docx_available = True
except ImportError:
    pass

try:
    import pypdf as _pypdf
    _pypdf_available = True
except ImportError:
    pass

# MIME type → extension map used for auto-detection when no mime_type is given.
_EXT_TO_MIME: dict[str, str] = {
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".pdf": "application/pdf",
    ".txt": "text/plain",
    ".md": "text/markdown",
}

_MIME_TO_EXT: dict[str, str] = {v: k for k, v in _EXT_TO_MIME.items()}

# Heading regex for markdown
_HEADING_RE: re.Pattern[str] = re.compile(r"^#{1,6}\s+(.+)", re.MULTILINE)


class DocumentParser:
    """Parse a file path into a :class:`~depthfusion.ingest.models.ParsedDocument`.

    Args:
        default_acl_allow:   Default ACL principal list applied when the
                             source connector does not supply one.
        default_classification: Default classification label.
        max_bytes:           Parse budget in bytes.  Documents larger than
                             this threshold are quarantined instead of parsed.
                             ``None`` (default) reads from the environment
                             variable ``DEPTHFUSION_PARSE_MAX_BYTES``.
                             ``0`` disables the budget entirely.
        quarantine_store:    Quarantine store instance to receive oversized
                             documents.  ``None`` uses a
                             :class:`FileQuarantineStore` writing to
                             ``data/quarantine/`` (or the module-level
                             :class:`~depthfusion.parsers.documents.base.QuarantineStore`
                             for backward compatibility when an explicit
                             in-memory store is preferred).  Pass an
                             explicit in-memory
                             :class:`~depthfusion.parsers.documents.base.QuarantineStore`
                             to avoid disk I/O in tests.
    """

    def __init__(
        self,
        default_acl_allow: list[str] | None = None,
        default_classification: str = "internal",
        max_bytes: int | None = None,
        quarantine_store=None,  # QuarantineStore | FileQuarantineStore | None
    ) -> None:
        self._default_acl = default_acl_allow or []
        self._default_classification = default_classification
        self._max_bytes: int = max_bytes if max_bytes is not None else _get_parse_max_bytes()
        if quarantine_store is None:
            from depthfusion.parsers.documents.base import get_quarantine_store as _get_qs
            self._quarantine_store = _get_qs()
        else:
            self._quarantine_store = quarantine_store

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def parse(
        self,
        path: str,
        mime_type: str | None = None,
        *,
        acl_allow: list[str] | None = None,
        classification: str | None = None,
    ) -> ParsedDocument:
        """Parse a document at *path* and return a :class:`ParsedDocument`.

        Args:
            path:           Absolute or relative file-system path.
            mime_type:      MIME type override.  If omitted, detected from
                            the file extension.
            acl_allow:      Per-document ACL override.
            classification: Per-document classification override.

        Returns:
            A populated :class:`ParsedDocument`.

        Raises:
            ValueError: If the file extension / MIME type is not supported,
                        or if the document exceeds the parse budget.
            FileNotFoundError: If *path* does not exist.
        """
        p = pathlib.Path(path)
        if not p.exists():
            raise FileNotFoundError(f"Document not found: {path}")

        resolved_mime = mime_type or _EXT_TO_MIME.get(p.suffix.lower(), "")
        if not resolved_mime:
            raise ValueError(
                f"Unsupported file extension '{p.suffix}' — "
                f"provide mime_type explicitly or use .docx/.pdf/.txt/.md"
            )

        raw_bytes = p.read_bytes()

        # T-599 — parse-budget enforcement
        if self._max_bytes > 0 and len(raw_bytes) > self._max_bytes:
            import datetime

            from depthfusion.parsers.documents.base import QuarantineEntry

            reason = (
                f"Document size {len(raw_bytes):,} bytes exceeds parse budget "
                f"{self._max_bytes:,} bytes"
            )
            entry = QuarantineEntry(
                source_id=str(p.resolve()),
                error_message=reason,
                timestamp=datetime.datetime.now(datetime.timezone.utc).isoformat(),
                raw_size_bytes=len(raw_bytes),
            )
            self._quarantine_store.add(entry)
            raise ValueError(reason)

        if resolved_mime in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ):
            text, metadata = self._parse_docx(raw_bytes, str(p))
        elif resolved_mime == "application/pdf":
            text, metadata = self._parse_pdf(raw_bytes, str(p))
        elif resolved_mime in ("text/plain", "text/markdown"):
            text, metadata = self._parse_text(raw_bytes, p.suffix.lower())
        else:
            raise ValueError(f"Unsupported MIME type: {resolved_mime}")

        metadata.setdefault("source_path", str(p))
        metadata.setdefault("filename", p.name)

        return ParsedDocument(
            source_id=str(p.resolve()),
            text=text,
            metadata=metadata,
            acl_allow=acl_allow if acl_allow is not None else list(self._default_acl),
            classification=classification or self._default_classification,
            mime_type=resolved_mime,
        )

    # ------------------------------------------------------------------
    # Format-specific helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _parse_docx(data: bytes, path: str) -> tuple[str, dict[str, str]]:
        """Extract text and metadata from a .docx file."""
        if not _docx_available:
            raise ImportError(
                "python-docx is required for .docx parsing. "
                "Install it with: pip install python-docx"
            )
        doc = _docx.Document(io.BytesIO(data))
        parts: list[str] = []
        for para in doc.paragraphs:
            stripped = para.text.strip()
            if stripped:
                parts.append(stripped)

        # Extract core properties for metadata
        metadata: dict[str, str] = {}
        cp = doc.core_properties
        if cp.title:
            metadata["title"] = str(cp.title)
        if cp.author:
            metadata["author"] = str(cp.author)
        if cp.created:
            metadata["created"] = str(cp.created)

        return "\n\n".join(parts), metadata

    @staticmethod
    def _parse_pdf(data: bytes, path: str) -> tuple[str, dict[str, str]]:
        """Extract text and metadata from a .pdf file."""
        if not _pypdf_available:
            raise ImportError(
                "pypdf is required for .pdf parsing. "
                "Install it with: pip install pypdf"
            )
        reader = _pypdf.PdfReader(io.BytesIO(data))
        parts: list[str] = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            stripped = page_text.strip()
            if stripped:
                parts.append(stripped)

        metadata: dict[str, str] = {}
        info = reader.metadata
        if info:
            if info.title:
                metadata["title"] = str(info.title)
            if info.author:
                metadata["author"] = str(info.author)
            if info.creation_date:
                metadata["created"] = str(info.creation_date)

        return "\n\n".join(parts), metadata

    @staticmethod
    def _parse_text(data: bytes, extension: str) -> tuple[str, dict[str, str]]:
        """Decode plain text or Markdown, extract a best-effort title."""
        try:
            text = data.decode("utf-8-sig")
        except UnicodeDecodeError:
            text = data.decode("latin-1")

        metadata: dict[str, str] = {}
        if extension == ".md":
            m = _HEADING_RE.search(text)
            if m:
                metadata["title"] = m.group(1).strip()
        if "title" not in metadata:
            for line in text.splitlines():
                stripped = line.strip()
                if stripped:
                    metadata["title"] = stripped[:100]
                    break

        return text, metadata


__all__ = ["DocumentParser", "FileQuarantineStore", "_DEFAULT_PARSE_MAX_BYTES"]
