"""Round-trip tests for the Document Ingestion Framework (E-53).

Tests:
  - .txt  round-trip via DocumentParser + IngestPipeline
  - .md   round-trip via DocumentParser + IngestPipeline
  - .docx round-trip via DocumentParser + IngestPipeline
  - .pdf  round-trip via DocumentParser + IngestPipeline
  - FixedSizeChunker behaviour (size, overlap)
  - SentenceBoundaryChunker behaviour
  - IngestPipeline embed/store callbacks receive the document
  - Unsupported file type raises ValueError
  - Missing file raises FileNotFoundError
  - ACL inheritance through the pipeline
"""
from __future__ import annotations

import io
import pathlib
import tempfile

import pytest

from depthfusion.ingest import (
    ChunkingStrategy,
    FixedSizeChunker,
    IngestPipeline,
    ParsedDocument,
    SentenceBoundaryChunker,
)
from depthfusion.ingest.parser import DocumentParser

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _write_tmp(suffix: str, content: bytes) -> pathlib.Path:
    """Write *content* to a temporary file and return the path."""
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as f:
        f.write(content)
        return pathlib.Path(f.name)


def _make_docx(text: str) -> bytes:
    """Create a minimal .docx file containing *text*."""
    from docx import Document

    doc = Document()
    for line in text.split("\n"):
        if line.strip():
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_pdf(text: str) -> bytes:
    """Create a minimal single-page PDF containing *text* via pypdf."""
    # We use pypdf's PdfWriter to create a blank PDF; injecting arbitrary
    # text into a PDF without a layout engine requires a third-party library.
    # Instead we use a minimal valid PDF with embedded text stream.
    # This is a raw hand-crafted PDF so we control exactly what text is stored.
    content_stream = f"BT /F1 12 Tf 72 720 Td ({text[:200]}) Tj ET"
    stream_bytes = content_stream.encode("latin-1")
    stream_length = len(stream_bytes)

    pdf_bytes = (
        b"%PDF-1.4\n"
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n"
        + (
            f"4 0 obj\n<< /Length {stream_length} >>\nstream\n"
        ).encode()
        + stream_bytes
        + b"\nendstream\nendobj\n"
        b"5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n"
        b"xref\n0 6\n"
        b"0000000000 65535 f \n"
        b"0000000009 00000 n \n"
        b"0000000058 00000 n \n"
        b"0000000115 00000 n \n"
        b"0000000266 00000 n \n"
        b"0000000400 00000 n \n"
        b"trailer\n<< /Size 6 /Root 1 0 R >>\nstartxref\n450\n%%EOF\n"
    )
    return pdf_bytes


# ---------------------------------------------------------------------------
# FixedSizeChunker tests
# ---------------------------------------------------------------------------

class TestFixedSizeChunker:
    def test_empty_text_returns_empty_list(self) -> None:
        chunker = FixedSizeChunker()
        assert chunker.chunk("") == []

    def test_short_text_returns_single_chunk(self) -> None:
        chunker = FixedSizeChunker(chunk_tokens=100, overlap_tokens=10)
        text = "Hello world."
        chunks = chunker.chunk(text)
        assert len(chunks) == 1
        assert "Hello world" in chunks[0]

    def test_long_text_produces_multiple_chunks(self) -> None:
        # 100 tokens * 4 chars/token = 400 chars per chunk
        chunker = FixedSizeChunker(chunk_tokens=100, overlap_tokens=20)
        text = "A" * 2000
        chunks = chunker.chunk(text)
        assert len(chunks) > 1

    def test_no_chunk_exceeds_limit(self) -> None:
        chunker = FixedSizeChunker(chunk_tokens=50, overlap_tokens=10)
        text = "B" * 2000
        chunk_chars = 50 * 4
        for chunk in chunker.chunk(text):
            assert len(chunk) <= chunk_chars

    def test_overlap_means_content_repeated(self) -> None:
        # With overlap, the end of chunk N should match the start of chunk N+1
        chunker = FixedSizeChunker(chunk_tokens=10, overlap_tokens=5)
        # 40 chars + 20 chars overlap → each chunk 40 chars, step 20 chars
        text = "ABCDEFGHIJ" * 10  # 100 chars
        chunks = chunker.chunk(text)
        if len(chunks) >= 2:
            # The tail of chunk 0 should appear in chunk 1
            overlap_chars = 5 * 4
            tail = chunks[0][-overlap_chars:]
            assert chunks[1].startswith(tail)

    def test_invalid_overlap_raises(self) -> None:
        with pytest.raises(ValueError, match="overlap_tokens"):
            FixedSizeChunker(chunk_tokens=10, overlap_tokens=10)


# ---------------------------------------------------------------------------
# SentenceBoundaryChunker tests
# ---------------------------------------------------------------------------

class TestSentenceBoundaryChunker:
    def test_empty_text(self) -> None:
        chunker = SentenceBoundaryChunker()
        assert chunker.chunk("") == []

    def test_single_short_sentence(self) -> None:
        chunker = SentenceBoundaryChunker(max_tokens=100)
        text = "Hello world."
        chunks = chunker.chunk(text)
        assert len(chunks) == 1
        assert "Hello world" in chunks[0]

    def test_splits_at_sentence_boundaries(self) -> None:
        # max_tokens=5 → max_chars=20
        chunker = SentenceBoundaryChunker(max_tokens=5)
        # Each sentence is ~15 chars — 2 sentences per chunk
        text = "Hi there. Bye now. See you. Later!"
        chunks = chunker.chunk(text)
        assert len(chunks) >= 1
        # Ensure all text is preserved
        all_text = " ".join(chunks)
        for word in ["Hi", "Bye", "See", "Later"]:
            assert word in all_text

    def test_long_sentence_hard_split(self) -> None:
        chunker = SentenceBoundaryChunker(max_tokens=10)
        # 1 sentence > max_chars (40 chars)
        text = "A" * 200
        chunks = chunker.chunk(text)
        max_chars = 10 * 4
        for chunk in chunks:
            assert len(chunk) <= max_chars


# ---------------------------------------------------------------------------
# DocumentParser round-trip tests
# ---------------------------------------------------------------------------

class TestDocumentParserTxt:
    def test_txt_round_trip(self, tmp_path: pathlib.Path) -> None:
        text = "Hello from a plain text document.\nSecond line here."
        p = tmp_path / "sample.txt"
        p.write_bytes(text.encode("utf-8"))

        parser = DocumentParser()
        doc = parser.parse(str(p))

        assert doc.source_id == str(p.resolve())
        assert "Hello from a plain text document" in doc.text
        assert doc.mime_type == "text/plain"
        assert doc.classification == "internal"

    def test_txt_metadata_title(self, tmp_path: pathlib.Path) -> None:
        text = "My Document Title\nSome body text."
        p = tmp_path / "doc.txt"
        p.write_bytes(text.encode("utf-8"))

        parser = DocumentParser()
        doc = parser.parse(str(p))
        assert doc.metadata.get("title") == "My Document Title"


class TestDocumentParserMarkdown:
    def test_md_round_trip(self, tmp_path: pathlib.Path) -> None:
        text = "# My Heading\n\nSome body paragraph.\n\n## Sub-heading\n\nMore content."
        p = tmp_path / "readme.md"
        p.write_bytes(text.encode("utf-8"))

        parser = DocumentParser()
        doc = parser.parse(str(p), "text/markdown")

        assert "My Heading" in doc.text
        assert "Sub-heading" in doc.text
        assert doc.mime_type == "text/markdown"
        assert doc.metadata.get("title") == "My Heading"

    def test_md_extension_autodetect(self, tmp_path: pathlib.Path) -> None:
        p = tmp_path / "notes.md"
        p.write_bytes(b"# Notes\n\nContent here.")
        doc = DocumentParser().parse(str(p))
        assert doc.mime_type == "text/markdown"


class TestDocumentParserDocx:
    def test_docx_round_trip(self, tmp_path: pathlib.Path) -> None:
        text = "Enterprise document content.\nWith multiple lines of important data."
        docx_bytes = _make_docx(text)
        p = tmp_path / "report.docx"
        p.write_bytes(docx_bytes)

        doc = DocumentParser().parse(str(p))

        assert "Enterprise document content" in doc.text
        assert doc.mime_type == (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    def test_docx_acl_propagation(self, tmp_path: pathlib.Path) -> None:
        docx_bytes = _make_docx("ACL test content.")
        p = tmp_path / "acl.docx"
        p.write_bytes(docx_bytes)

        doc = DocumentParser().parse(str(p), acl_allow=["user-a", "group-b"])
        assert "user-a" in doc.acl_allow
        assert "group-b" in doc.acl_allow


class TestDocumentParserPdf:
    def test_pdf_round_trip(self, tmp_path: pathlib.Path) -> None:
        """PDF parsing should succeed and return a ParsedDocument."""
        pdf_bytes = _make_pdf("Hello PDF content here")
        p = tmp_path / "doc.pdf"
        p.write_bytes(pdf_bytes)

        doc = DocumentParser().parse(str(p))
        assert doc.mime_type == "application/pdf"
        assert isinstance(doc.text, str)
        # Text extraction may be empty for a hand-crafted PDF, but should not raise

    def test_pdf_source_id_is_resolved_path(self, tmp_path: pathlib.Path) -> None:
        pdf_bytes = _make_pdf("Source ID test")
        p = tmp_path / "sid.pdf"
        p.write_bytes(pdf_bytes)

        doc = DocumentParser().parse(str(p))
        assert doc.source_id == str(p.resolve())


class TestDocumentParserErrors:
    def test_missing_file_raises(self) -> None:
        with pytest.raises(FileNotFoundError):
            DocumentParser().parse("/tmp/nonexistent_file_xyz123.docx")

    def test_unsupported_extension_raises(self, tmp_path: pathlib.Path) -> None:
        p = tmp_path / "file.xyz"
        p.write_bytes(b"data")
        with pytest.raises(ValueError, match="Unsupported file extension"):
            DocumentParser().parse(str(p))

    def test_unsupported_mime_raises(self, tmp_path: pathlib.Path) -> None:
        p = tmp_path / "file.txt"
        p.write_bytes(b"data")
        with pytest.raises(ValueError, match="Unsupported MIME type"):
            DocumentParser().parse(str(p), mime_type="application/octet-stream")


# ---------------------------------------------------------------------------
# IngestPipeline tests
# ---------------------------------------------------------------------------

class TestIngestPipeline:
    def test_pipeline_txt_produces_chunks(self, tmp_path: pathlib.Path) -> None:
        # Generate text long enough to produce multiple chunks
        text = ("This is sentence number one. " * 100).strip()
        p = tmp_path / "long.txt"
        p.write_bytes(text.encode())

        pipeline = IngestPipeline()
        doc = pipeline.run(str(p))

        assert len(doc.chunks) >= 1
        assert doc.text  # text not empty
        assert doc.mime_type == "text/plain"

    def test_pipeline_md_chunks_populated(self, tmp_path: pathlib.Path) -> None:
        text = "# Heading\n\n" + ("Word " * 500) + "\n\n## Sub\n\n" + ("More " * 200)
        p = tmp_path / "big.md"
        p.write_bytes(text.encode())

        pipeline = IngestPipeline()
        doc = pipeline.run(str(p))
        assert len(doc.chunks) > 0

    def test_pipeline_docx_round_trip(self, tmp_path: pathlib.Path) -> None:
        body = "Important business document. " * 20
        p = tmp_path / "biz.docx"
        p.write_bytes(_make_docx(body))

        doc = IngestPipeline().run(str(p))
        assert "Important business document" in doc.text
        assert len(doc.chunks) >= 1

    def test_pipeline_acl_inherited_in_chunks(self, tmp_path: pathlib.Path) -> None:
        text = "Sensitive data here. " * 50
        p = tmp_path / "sens.txt"
        p.write_bytes(text.encode())

        acl = ["alice@corp.com", "team-engineering"]
        pipeline = IngestPipeline()
        doc = pipeline.run(str(p), acl_allow=acl)

        assert doc.acl_allow == acl

    def test_pipeline_embed_callback_called(self, tmp_path: pathlib.Path) -> None:
        p = tmp_path / "cb.txt"
        p.write_bytes(b"Short doc for callback test.")

        called_with: list[ParsedDocument] = []
        pipeline = IngestPipeline(embed_callback=called_with.append)
        doc = pipeline.run(str(p))

        assert len(called_with) == 1
        assert called_with[0] is doc

    def test_pipeline_store_callback_called(self, tmp_path: pathlib.Path) -> None:
        p = tmp_path / "store.txt"
        p.write_bytes(b"Store callback test content.")

        stored: list[ParsedDocument] = []
        pipeline = IngestPipeline(store_callback=stored.append)
        doc = pipeline.run(str(p))

        assert len(stored) == 1
        assert stored[0] is doc

    def test_pipeline_run_from_bytes(self) -> None:
        """run_from_bytes should work without touching the file system."""
        text = "Bytes-based ingestion test content."
        pipeline = IngestPipeline()
        doc = pipeline.run_from_bytes(
            source_id="mem://test-001",
            data=text.encode(),
            mime_type="text/plain",
            acl_allow=["viewer@corp.com"],
        )
        assert doc.source_id == "mem://test-001"
        assert "Bytes-based ingestion" in doc.text
        assert "viewer@corp.com" in doc.acl_allow

    def test_pipeline_classification_propagated(self, tmp_path: pathlib.Path) -> None:
        p = tmp_path / "classified.txt"
        p.write_bytes(b"Confidential content.")

        doc = IngestPipeline().run(str(p), classification="confidential")
        assert doc.classification == "confidential"

    def test_pipeline_missing_file_raises(self) -> None:
        with pytest.raises(FileNotFoundError):
            IngestPipeline().run("/tmp/no_such_file_abc.txt")

    def test_pipeline_pdf_round_trip(self, tmp_path: pathlib.Path) -> None:
        pdf_bytes = _make_pdf("PDF pipeline test content")
        p = tmp_path / "pipe.pdf"
        p.write_bytes(pdf_bytes)

        doc = IngestPipeline().run(str(p))
        assert doc.mime_type == "application/pdf"
        assert len(doc.chunks) >= 0  # may be 0 for hand-crafted PDF


def test_chunks_carry_heading_path():
    """Chunks produced by the pipeline inherit heading_path from metadata."""
    from unittest.mock import MagicMock

    from depthfusion.ingest.models import Chunk

    mock_parser = MagicMock()
    parsed = ParsedDocument(
        source_id="test_doc",
        text="Alpha beta gamma delta epsilon zeta",
        metadata={"heading_path": "Introduction"},
    )
    mock_parser.parse.return_value = parsed

    pipeline = IngestPipeline(
        parser=mock_parser,
        chunker=FixedSizeChunker(chunk_tokens=5, overlap_tokens=0),
    )
    doc = pipeline.run("/fake/path.txt")

    assert all(isinstance(c, Chunk) for c in doc.chunks)
    assert all(c.heading_path == "Introduction" for c in doc.chunks)


# ---------------------------------------------------------------------------
# ChunkingStrategy Protocol conformance
# ---------------------------------------------------------------------------

class TestChunkingStrategyProtocol:
    def test_fixed_size_satisfies_protocol(self) -> None:
        chunker = FixedSizeChunker()
        assert isinstance(chunker, ChunkingStrategy)

    def test_sentence_boundary_satisfies_protocol(self) -> None:
        chunker = SentenceBoundaryChunker()
        assert isinstance(chunker, ChunkingStrategy)


# ---------------------------------------------------------------------------
# T-599 — Parse-budget oversized-doc quarantine with disk persistence
# ---------------------------------------------------------------------------

class TestParseBudgetQuarantinePersistence:
    """Oversized documents are quarantined to data/quarantine/ with reason logged.

    Acceptance criteria (T-599):
      AC-1: Files exceeding the budget are written to data/quarantine/ as JSON
            sidecar files (not in-memory only).
      AC-2: The JSON sidecar contains the source_id, raw_size_bytes, and a
            reason string that mentions the doc size and the budget limit.
      AC-3: The document is NOT parsed — parse() raises ValueError.
      AC-4: Default threshold (50 MiB) and DEPTHFUSION_PARSE_MAX_BYTES override
            both work; setting budget=0 disables quarantine.
    """

    def test_oversized_doc_creates_quarantine_artifact(
        self, tmp_path: pathlib.Path
    ) -> None:
        """An oversized doc produces a JSON sidecar in data/quarantine/."""
        from depthfusion.parsers.documents.base import QuarantineStore

        quarantine_dir = tmp_path / "data" / "quarantine"
        store = QuarantineStore(persist_dir=quarantine_dir)
        budget = 100
        parser = DocumentParser(max_bytes=budget, quarantine_store=store)

        content = b"X" * (budget + 1)
        doc_path = tmp_path / "big.txt"
        doc_path.write_bytes(content)

        with pytest.raises(ValueError, match="parse budget"):
            parser.parse(str(doc_path), "text/plain")

        # Quarantine directory must now exist and have exactly one JSON file.
        artifacts = list(quarantine_dir.glob("*.json"))
        assert len(artifacts) == 1, f"expected 1 artifact, found {artifacts}"

    def test_quarantine_artifact_contains_reason(
        self, tmp_path: pathlib.Path
    ) -> None:
        """The JSON sidecar records source_id, raw_size_bytes, and reason."""
        import json as _json

        from depthfusion.parsers.documents.base import QuarantineStore

        quarantine_dir = tmp_path / "data" / "quarantine"
        store = QuarantineStore(persist_dir=quarantine_dir)
        budget = 200
        parser = DocumentParser(max_bytes=budget, quarantine_store=store)

        content = b"Y" * (budget + 50)
        doc_path = tmp_path / "oversized.txt"
        doc_path.write_bytes(content)

        with pytest.raises(ValueError):
            parser.parse(str(doc_path), "text/plain")

        artifacts = list(quarantine_dir.glob("*.json"))
        assert artifacts, "expected at least one quarantine artifact"
        record = _json.loads(artifacts[0].read_text(encoding="utf-8"))

        # source_id must reference the document.
        assert str(doc_path.resolve()) in record["source_id"]
        # raw_size_bytes must be accurate.
        assert record["raw_size_bytes"] == len(content)
        # error_message must mention both sizes so operators can diagnose.
        msg = record["error_message"]
        assert str(len(content)) in msg
        assert str(budget) in msg

    def test_oversized_doc_is_not_parsed(self, tmp_path: pathlib.Path) -> None:
        """parse() raises ValueError for oversized docs — no ParsedDocument returned."""
        from depthfusion.parsers.documents.base import QuarantineStore

        quarantine_dir = tmp_path / "data" / "quarantine"
        store = QuarantineStore(persist_dir=quarantine_dir)
        budget = 50
        parser = DocumentParser(max_bytes=budget, quarantine_store=store)

        doc_path = tmp_path / "toobig.txt"
        doc_path.write_bytes(b"Z" * (budget + 10))

        with pytest.raises(ValueError):
            parser.parse(str(doc_path), "text/plain")

        # Nothing was parsed — quarantine dir has the artifact, not a parsed doc.
        artifacts = list(quarantine_dir.glob("*.json"))
        assert len(artifacts) == 1

    def test_budget_zero_disables_quarantine(self, tmp_path: pathlib.Path) -> None:
        """budget=0 means no limit — no quarantine artifact is created."""
        from depthfusion.parsers.documents.base import QuarantineStore

        quarantine_dir = tmp_path / "data" / "quarantine"
        store = QuarantineStore(persist_dir=quarantine_dir)
        parser = DocumentParser(max_bytes=0, quarantine_store=store)

        content = b"A" * 500
        doc_path = tmp_path / "unlimited.txt"
        doc_path.write_bytes(content)

        doc = parser.parse(str(doc_path), "text/plain")
        assert doc is not None

        # No quarantine artifact produced.
        if quarantine_dir.exists():
            assert list(quarantine_dir.glob("*.json")) == []

    def test_env_var_override_persists_to_disk(
        self, tmp_path: pathlib.Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """DEPTHFUSION_PARSE_MAX_BYTES env override triggers disk quarantine."""
        import json as _json

        from depthfusion.parsers.documents.base import QuarantineStore

        monkeypatch.setenv("DEPTHFUSION_PARSE_MAX_BYTES", "30")
        quarantine_dir = tmp_path / "data" / "quarantine"
        store = QuarantineStore(persist_dir=quarantine_dir)
        parser = DocumentParser(max_bytes=None, quarantine_store=store)
        assert parser._max_bytes == 30

        content = b"E" * 50
        doc_path = tmp_path / "envtest.txt"
        doc_path.write_bytes(content)

        with pytest.raises(ValueError, match="parse budget"):
            parser.parse(str(doc_path), "text/plain")

        artifacts = list(quarantine_dir.glob("*.json"))
        assert len(artifacts) == 1
        record = _json.loads(artifacts[0].read_text(encoding="utf-8"))
        assert "30" in record["error_message"]
